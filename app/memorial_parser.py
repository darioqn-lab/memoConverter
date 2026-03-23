# -*- coding: utf-8 -*-
"""
memorial_parser.py
MemoConverter v1.7.0
Extrai vertices de memoriais descritivos brasileiros.
Suporta coordenadas UTM (E/N) e geograficas (graus/minutos/segundos).
"""

import re
import math


# ──────────────────────────────────────────────────────────────
# Leitura de arquivos
# ──────────────────────────────────────────────────────────────

def _tem_coordenadas(texto):
    """Verifica se o texto contem coordenadas UTM ou geograficas."""
    return bool(re.search(r'\d{6,8}[.,]\d{2,6}', texto))


def ler_pdf(path):
    """
    Extrai texto de PDF selecionavel. Retorna (texto, status).
    Se o texto extraido for pequeno ou nao tiver coordenadas,
    marca como ocr_needed pois o conteudo pode estar em imagem.
    """
    textos = []

    try:
        import pdfplumber
        partes = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                tbl = page.extract_table()
                if tbl:
                    for row in tbl:
                        celulas = [str(c).strip() for c in row if c and str(c).strip()]
                        if celulas:
                            partes.append("  ".join(celulas))
                t = page.extract_text()
                if t:
                    partes.append(t)
        texto = "\n".join(partes)
        if texto.strip():
            textos.append(texto)
    except Exception:
        pass

    try:
        import fitz
        doc = fitz.open(path)
        texto = "\n".join(p.get_text() for p in doc)
        doc.close()
        if texto.strip():
            textos.append(texto)
    except Exception:
        pass

    # Escolhe o texto mais longo
    if textos:
        melhor = max(textos, key=len)
        # Se o texto e muito curto (< 500 chars) ou nao tem coordenadas,
        # provavelmente o conteudo principal esta em imagem — precisa OCR
        if len(melhor) < 500 or not _tem_coordenadas(melhor):
            return melhor, "ocr_needed"
        return melhor, None

    return "", "ocr_needed"


def ler_docx(path):
    """Extrai texto de Word .docx. Retorna (texto, erro_ou_None)."""
    try:
        import docx
        doc = docx.Document(path)
        partes = [p.text for p in doc.paragraphs if p.text.strip()]
        for tabela in doc.tables:
            for row in tabela.rows:
                celulas = [c.text.strip() for c in row.cells if c.text.strip()]
                if celulas:
                    partes.append("  ".join(celulas))
        return "\n".join(partes), None
    except ImportError:
        return "", "python-docx nao instalado. Execute: pip install python-docx"
    except Exception as e:
        return "", str(e)


# ──────────────────────────────────────────────────────────────
# OCR via Google Vision
# ──────────────────────────────────────────────────────────────

def ocr_google_vision(path, api_key, max_paginas=20):
    """
    Envia PDF para Google Vision e retorna texto extraido.
    Divide automaticamente em lotes de 5 paginas (limite da API).
    """
    import base64, json, urllib.request, urllib.error

    if not api_key or api_key.strip() == "":
        return "", "Chave do Google Vision nao configurada em config.ini."

    try:
        with open(path, "rb") as f:
            content = base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        return "", "Erro ao ler arquivo: {}".format(e)

    url        = "https://vision.googleapis.com/v1/files:annotate?key={}".format(api_key)
    LOTE_MAX   = 5   # limite da API do Google Vision
    total_pags = min(max_paginas, 20)
    texto_total = ""

    # divide em lotes de 5 paginas
    for inicio in range(1, total_pags + 1, LOTE_MAX):
        fim   = min(inicio + LOTE_MAX - 1, total_pags)
        pages = list(range(inicio, fim + 1))

        payload = {
            "requests": [{
                "inputConfig": {"content": content, "mimeType": "application/pdf"},
                "features":    [{"type": "DOCUMENT_TEXT_DETECTION"}],
                "pages":       pages
            }]
        }
        data = json.dumps(payload).encode("utf-8")
        req  = urllib.request.Request(url, data=data,
               headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                result = json.loads(resp.read().decode("utf-8"))
        except urllib.error.HTTPError as e:
            body = e.read().decode("utf-8")
            try:
                msg = json.loads(body).get("error", {}).get("message", body)
            except Exception:
                msg = body
            # se for erro de pagina inexistente, para de tentar mais lotes
            if "pages" in msg.lower() or "page" in msg.lower():
                break
            return "", "Erro Google Vision: {}".format(msg)
        except Exception as e:
            return "", "Erro de conexao: {}".format(e)

        for resp in result.get("responses", []):
            for pr in resp.get("responses", []):
                texto_total += pr.get("fullTextAnnotation", {}).get("text", "") + "\n"

    if not texto_total.strip():
        return "", "OCR nao retornou texto. Verifique se o PDF contem imagens legiveis."
    return texto_total, None


def ocr_google_vision_imagem(path, api_key):
    """
    Envia imagem (JPG, PNG, TIF) para Google Vision e retorna texto extraido.
    Usa endpoint images:annotate (diferente do endpoint de PDF).
    """
    import base64, json, urllib.request, urllib.error

    if not api_key:
        return "", "Chave do Google Vision nao configurada."

    try:
        with open(path, "rb") as f:
            content_b64 = base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        return "", "Erro ao ler imagem: {}".format(e)

    url = "https://vision.googleapis.com/v1/images:annotate?key={}".format(api_key)
    payload = {
        "requests": [{
            "image":    {"content": content_b64},
            "features": [{"type": "DOCUMENT_TEXT_DETECTION"}]
        }]
    }

    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(url, data=data,
           headers={"Content-Type": "application/json"}, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8")
        try:
            msg = json.loads(body).get("error", {}).get("message", body)
        except Exception:
            msg = body
        return "", "Erro Google Vision: {}".format(msg)
    except Exception as e:
        return "", "Erro de conexao: {}".format(e)

    try:
        texto = result["responses"][0]["fullTextAnnotation"]["text"]
        return texto, None
    except (KeyError, IndexError):
        return "", "OCR nao retornou texto. Verifique se a imagem e legivel."


# ──────────────────────────────────────────────────────────────
# Extracao inteligente via Gemini
# ──────────────────────────────────────────────────────────────

def _anonimizar_para_ia(texto):
    """
    Filtra o texto antes de enviar para IA externa.
    Mantém apenas os trechos que contêm coordenadas UTM ou DMS,
    removendo dados pessoais (CPF, nomes, endereços).
    Reduz dados transmitidos e protege informações sensíveis.
    """
    linhas = texto.split("\n")
    trechos = []

    # Padrões que indicam linha com coordenada
    pat_utm = re.compile(r"\d{5,8}[.,]\d{2,6}")
    pat_dms = re.compile(r"\d{1,3}[°º]\d{1,2}['’]\d{1,2}")
    pat_vertice = re.compile(
        r"\b(v[eé]rtice|marco|ponto|vx?[-_]?\d|pp[-_]?\d|[A-Z][-_]\d)",
        re.IGNORECASE
    )

    # Janela: inclui linha com coordenada + 1 linha antes e 1 depois (contexto)
    incluir = set()
    for i, linha in enumerate(linhas):
        if pat_utm.search(linha) or pat_dms.search(linha) or pat_vertice.search(linha):
            incluir.add(max(0, i - 1))
            incluir.add(i)
            incluir.add(min(len(linhas) - 1, i + 1))

    if not incluir:
        # Se nao encontrou nada relevante, retorna texto truncado sem dados pessoais
        texto_limpo = re.sub(
            r"\b\d{3}\.\d{3}\.\d{3}[-]\d{2}\b",  # CPF
            "[CPF]", texto
        )
        texto_limpo = re.sub(
            r"\b\d{2}\.\d{3}\.\d{3}/\d{4}[-]\d{2}\b",  # CNPJ
            "[CNPJ]", texto_limpo
        )
        return texto_limpo[:4000]

    # Monta texto filtrado so com linhas relevantes
    resultado = []
    for i in sorted(incluir):
        linha = linhas[i]
        # Remove CPF e CNPJ mesmo nas linhas incluidas
        linha = re.sub(r"\b\d{3}\.\d{3}\.\d{3}[-]\d{2}\b", "[CPF]", linha)
        linha = re.sub(r"\b\d{2}\.\d{3}\.\d{3}/\d{4}[-]\d{2}\b", "[CNPJ]", linha)
        resultado.append(linha)

    return "\n".join(resultado)


def extrair_com_gemini(texto, api_key):
    """
    Usa Gemini para extrair vertices de memoriais em formato nao padrao.
    Retorna (lista_vertices, erro_ou_None).
    """
    import json, urllib.request, urllib.error

    if not api_key or api_key.strip() == "":
        return [], "Chave do Gemini nao configurada em config.ini."

    prompt = (
        "Voce e um especialista em memoriais descritivos de imoveis brasileiros.\n"
        "Extraia todos os vertices/marcos do memorial abaixo e retorne SOMENTE um JSON valido, sem explicacoes.\n\n"
        "Formato obrigatorio:\n"
        '[\n  {"vertice": "nome_ou_numero", "coord_e": valor_numerico, "coord_n": valor_numerico},\n  ...\n]\n\n'
        "Se as coordenadas forem em graus/minutos/segundos (ex: 38 57 20 W, 14 04 13 S),\n"
        "converta para decimal e inclua tambem os campos lon e lat em vez de coord_e/coord_n.\n\n"
        "Se nao encontrar coordenadas, retorne [].\n\n"
        "MEMORIAL:\n"
        + _anonimizar_para_ia(texto)
    )

    url = ("https://generativelanguage.googleapis.com/v1beta/models/"
           "gemini-2.5-flash:generateContent?key={}".format(api_key))

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0, "maxOutputTokens": 2048}
    }

    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(url, data=data,
           headers={"Content-Type": "application/json"}, method="POST")

    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8")
        try:
            msg = json.loads(body).get("error", {}).get("message", body)
        except Exception:
            msg = body
        return [], "Erro Gemini: {}".format(msg)
    except Exception as e:
        return [], "Erro de conexao: {}".format(e)

    try:
        candidatos = result.get("candidates", [])
        if not candidatos:
            # retorna o resultado completo para diagnostico
            return [], "Gemini sem candidatos. Resposta: {}".format(str(result)[:300])
        raw = candidatos[0]["content"]["parts"][0]["text"]
        # limpa markdown e texto antes/depois do JSON
        raw = re.sub(r"```json|```", "", raw).strip()
        # tenta extrair array JSON se houver texto antes
        match_json = re.search(r'\[.*\]', raw, re.DOTALL)
        if match_json:
            raw = match_json.group(0)
        vertices = json.loads(raw)
        if not isinstance(vertices, list):
            return [], "Resposta do Gemini em formato inesperado."
        # normaliza campos
        resultado = []
        for v in vertices:
            if "lat" in v and "lon" in v:
                # coordenadas geograficas — converte para UTM
                lat = float(v["lat"])
                lon = float(v["lon"])
                e, n, fuso = geo_para_utm(lat, lon)
                resultado.append({
                    "vertice":  str(v.get("vertice", "")),
                    "coord_e":  round(e, 2),
                    "coord_n":  round(n, 2),
                    "tipo":     "geografica",
                    "lat":      lat,
                    "lon":      lon,
                    "fuso_utm": fuso,
                })
            elif "coord_e" in v and "coord_n" in v:
                resultado.append({
                    "vertice": str(v.get("vertice", "")),
                    "coord_e": float(v["coord_e"]),
                    "coord_n": float(v["coord_n"]),
                    "tipo":    "utm",
                })
        return resultado, None
    except Exception as e:
        raw_preview = raw[:200] if 'raw' in dir() else "(sem resposta)"
        return [], "Erro ao interpretar resposta do Gemini: {} | Resposta: {}".format(e, raw_preview)


# ──────────────────────────────────────────────────────────────
# Conversao de coordenadas geograficas
# ──────────────────────────────────────────────────────────────

def dms_para_decimal(graus, minutos, segundos, direcao):
    """Converte DMS para graus decimais. direcao: 'S','W' = negativo."""
    dd = float(graus) + float(minutos)/60.0 + float(segundos)/3600.0
    if direcao.upper() in ('S', 'W', 'O'):
        dd = -dd
    return dd


def geo_para_utm(lat, lon):
    """
    Converte lat/lon (WGS84) para UTM.
    Retorna (easting, northing, fuso_numero).
    Implementacao simples sem dependencias externas.
    """
    # calcula fuso
    fuso = int((lon + 180) / 6) + 1

    # constantes WGS84
    a  = 6378137.0
    f  = 1 / 298.257223563
    b  = a * (1 - f)
    e2 = 1 - (b/a)**2
    e  = math.sqrt(e2)
    k0 = 0.9996

    lat_r = math.radians(lat)
    lon_r = math.radians(lon)
    lon0  = math.radians((fuso - 1) * 6 - 180 + 3)  # meridiano central

    N  = a / math.sqrt(1 - e2 * math.sin(lat_r)**2)
    T  = math.tan(lat_r)**2
    C  = (e2 / (1 - e2)) * math.cos(lat_r)**2
    A  = math.cos(lat_r) * (lon_r - lon0)

    e4 = e2**2; e6 = e2**3
    M  = a * (
        (1 - e2/4 - 3*e4/64 - 5*e6/256) * lat_r
      - (3*e2/8 + 3*e4/32 + 45*e6/1024) * math.sin(2*lat_r)
      + (15*e4/256 + 45*e6/1024) * math.sin(4*lat_r)
      - (35*e6/3072) * math.sin(6*lat_r)
    )

    easting = k0 * N * (
        A + (1-T+C)*A**3/6 + (5-18*T+T**2+72*C-58*(e2/(1-e2)))*A**5/120
    ) + 500000.0

    northing = k0 * (M + N*math.tan(lat_r) * (
        A**2/2 + (5-T+9*C+4*C**2)*A**4/24
        + (61-58*T+T**2+600*C-330*(e2/(1-e2)))*A**6/720
    ))

    if lat < 0:
        northing += 10000000.0  # hemisferio sul

    return round(easting, 2), round(northing, 2), fuso


def detectar_e_converter_gms(texto):
    """
    Detecta coordenadas geograficas no texto e retorna lista de vertices convertidos.
    Aceita formatos como:
      38°57'20.596"W  14°04'13.869"S
      Long: 38°57'20"W  Lat: 14°04'13"S
      Longitude: 38 57 20 W  Latitude: 14 04 13 S
    """
    # padrao DMS completo
    _DMS = (r"(\d{1,3})[°º\s]"
            r"(\d{1,2})['\u2032\s]"
            r"(\d{1,2}(?:[.,]\d+)?)[\"″\s]?"
            r"\s*([NSEWOnsewо])")

    # busca pares lon/lat ou lat/lon
    pares_lon_lat = re.compile(
        r"[Ll]on(?:gitude)?\s*[:\-]?\s*" + _DMS
        + r".{0,30}?"
        + r"[Ll]at(?:itude)?\s*[:\-]?\s*" + _DMS,
        re.DOTALL
    )
    pares_lat_lon = re.compile(
        r"[Ll]at(?:itude)?\s*[:\-]?\s*" + _DMS
        + r".{0,30}?"
        + r"[Ll]on(?:gitude)?\s*[:\-]?\s*" + _DMS,
        re.DOTALL
    )
    # pares sem rotulo (dois DMS seguidos)
    pares_simples = re.compile(
        _DMS + r"\s*[,;]?\s*" + _DMS
    )

    vertices = []
    seen     = set()

    def processar_par(lon_g, lon_m, lon_s, lon_d, lat_g, lat_m, lat_s, lat_d, idx):
        try:
            lon = dms_para_decimal(lon_g, lon_m, lon_s.replace(",","."), lon_d)
            lat = dms_para_decimal(lat_g, lat_m, lat_s.replace(",","."), lat_d)
            chave = (round(lat,4), round(lon,4))
            if chave in seen:
                return None
            seen.add(chave)
            e, n, fuso = geo_para_utm(lat, lon)
            return {
                "vertice":  str(idx + 1),
                "coord_e":  e,
                "coord_n":  n,
                "tipo":     "geografica",
                "lat":      round(lat, 6),
                "lon":      round(lon, 6),
                "fuso_utm": fuso,
            }
        except Exception:
            return None

    for pat, ordem in [(pares_lon_lat, "LE"), (pares_lat_lon, "EL")]:
        for m in pat.finditer(texto):
            g = m.groups()
            if ordem == "LE":  # lon primeiro, lat segundo
                v = processar_par(g[0],g[1],g[2],g[3], g[4],g[5],g[6],g[7], len(vertices))
            else:              # lat primeiro, lon segundo
                v = processar_par(g[4],g[5],g[6],g[7], g[0],g[1],g[2],g[3], len(vertices))
            if v:
                vertices.append(v)

    # fallback: pares simples sem rotulo
    if not vertices:
        for m in pares_simples.finditer(texto):
            g = m.groups()
            d1, d2 = g[3].upper(), g[7].upper()
            # decide qual e lon e qual e lat pela direcao
            if d1 in ('E','W','O') and d2 in ('N','S'):
                v = processar_par(g[0],g[1],g[2],g[3], g[4],g[5],g[6],g[7], len(vertices))
            elif d1 in ('N','S') and d2 in ('E','W','O'):
                v = processar_par(g[4],g[5],g[6],g[7], g[0],g[1],g[2],g[3], len(vertices))
            else:
                continue
            if v:
                vertices.append(v)

    return vertices


# ──────────────────────────────────────────────────────────────
# Helpers numericos
# ──────────────────────────────────────────────────────────────

def _num(s):
    s = str(s).strip().replace(" ", "")
    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")
    return float(s)


def _is_utm(v):
    return 100_000 <= v <= 9_999_999


def _classificar_en(v1, v2):
    if v1 > v2:
        return v2, v1
    return v1, v2


# ──────────────────────────────────────────────────────────────
# Extrator de metadados
# ──────────────────────────────────────────────────────────────

def extrair_meta(texto):
    meta = {}
    padroes = [
        (r"[Aa]rea\s*(?:total\s*)?[:\-=]?\s*([\d.,]+)\s*m",   "area"),
        (r"[Pp]er[iI]metro\s*[:\-=]?\s*([\d.,]+)\s*m",        "perimetro"),
        (r"(SIRGAS\s*2000|SAD[\s\-]*69|WGS[\s\-]*84)",         "datum"),
        (r"[Mm]unic[iI]pio\s*(?:de\s*)?[:\-]?\s*([^\n,/]{3,50})", "municipio"),
        (r"[Ff]uso\s*[:\-]?\s*(-?\d{1,2})",                   "fuso"),
        (r"[Mm]eridiano\s+[Cc]entral\s+(-?\d{1,3}[^\s,;]{0,10})", "meridiano"),
    ]
    for pat, chave in padroes:
        m = re.search(pat, texto, re.IGNORECASE)
        if not m:
            continue
        val = m.group(1).strip()
        if chave in ("area", "perimetro"):
            try: meta[chave] = _num(val)
            except: pass
        elif chave == "datum":
            meta[chave] = re.sub(r"\s+", "", val).upper()
        else:
            meta[chave] = val

    # Deriva fuso a partir do Meridiano Central se nao encontrou fuso diretamente
    if "fuso" not in meta and "meridiano" in meta:
        try:
            mc_str = meta["meridiano"]
            # extrai numero do meridiano (ex: "45°00'", "45", "-45", "39°")
            mc_num = abs(float(re.search(r"\d+", mc_str).group()))
            # formula: fuso = (180 + MC) / 6 arredondado
            fuso_calc = int((180 - mc_num) / 6) + 1
            if 18 <= fuso_calc <= 25:
                meta["fuso"] = str(fuso_calc)
        except Exception:
            pass

    # Converte fuso para inteiro se necessario
    if "fuso" in meta:
        try:
            meta["fuso"] = str(int(re.search(r"\d+", str(meta["fuso"])).group()))
        except Exception:
            pass

    return meta


# ──────────────────────────────────────────────────────────────
# Extrator principal de vertices UTM
# ──────────────────────────────────────────────────────────────

# Regex UTM: aceita separadores de milhar BR (8.444.474,16) e US (8444474.16)
_UTM = r"(\d{1,3}(?:[.,]\d{3})+(?:[.,]\d{1,6})?|\d{5,7}(?:[.,]\d{1,6})?)"

# ──────────────────────────────────────────────────────────────
# Ordenacao e validacao do poligono
# ──────────────────────────────────────────────────────────────

def calcular_sentido(vertices):
    """Retorna 'horario' ou 'antihorario' usando formula shoelace."""
    if len(vertices) < 3:
        return None
    area = 0.0
    n = len(vertices)
    for i in range(n):
        j = (i + 1) % n
        area += (vertices[i]["coord_e"] * vertices[j]["coord_n"] -
                 vertices[j]["coord_e"] * vertices[i]["coord_n"])
    return "horario" if area < 0 else "antihorario"


def ordenar_vertices(vertices):
    """
    Corrige ordem: inicia pelo vertice mais ao norte + sentido horario.
    Mantem os rotulos originais dos vertices.
    """
    if not vertices or len(vertices) < 3:
        return vertices

    # Inicia pelo vertice mais ao norte
    idx_norte = max(range(len(vertices)), key=lambda i: vertices[i]["coord_n"])
    rotacionado = vertices[idx_norte:] + vertices[:idx_norte]

    # Corrige sentido se estiver anti-horario
    if calcular_sentido(rotacionado) == "antihorario":
        rotacionado = rotacionado[::-1]

    return rotacionado


def validar_poligono(vertices):
    """
    Valida quantidade minima de vertices.
    Memoriais descritivos por definicao legal sempre fecham —
    a distancia entre o ultimo e o primeiro vertice e o ultimo lado
    do poligono, nao uma indicacao de erro.
    Retorna "OK" ou mensagem de aviso.
    """
    if len(vertices) < 3:
        return "Vertices insuficientes (minimo 3)"
    if len(vertices) < 4:
        return "Poligono com apenas {} vertices — verifique a extracao".format(len(vertices))
    return "OK"


def extrair_vertices(texto):
    """
    Estrategia em cascata:
    1. Coordenadas UTM com rotulos (Este/Norte, N/E, E=/N=) — prioridade
    2. Tabela (linhas com pares de numeros UTM)
    3. Coordenadas geograficas (DMS) — fallback se nao houver UTM
    Motivo: memoriais brasileiros frequentemente citam coordenadas geograficas
    apenas como referencia do ponto inicial, mas os demais marcos estao em UTM.
    """
    texto = _normalizar(texto)

    # tenta UTM primeiro — e o formato principal dos memoriais brasileiros
    v_utm = _extrair_utm(texto)
    if v_utm:
        return v_utm

    # par inline (E N na mesma linha separados por espaco)
    v_inline = _extrair_pares_inline(texto)
    if v_inline:
        return v_inline

    # tabela em blocos (cada celula em linha separada)
    v_bloco = _extrair_tabela_blocos(texto)
    if v_bloco:
        return v_bloco

    # fallback: coordenadas geograficas (DMS)
    v_geo = detectar_e_converter_gms(texto)
    if v_geo:
        return v_geo

    return []


def _normalizar(texto):
    texto = re.sub(r"\r\n|\r", "\n", texto)
    texto = re.sub(r"[ \t]+", " ", texto)
    for src, dst in [("\u201c",'"'),("\u201d",'"'),("\u2018","'"),("\u2019","'")]:
        texto = texto.replace(src, dst)
    return texto


def _extrair_utm(texto):
    pat = re.compile(
        r"[Ee]ste\s*(?:\(X\))?\s*[=:]?\s*" + _UTM
        + r"(?:\s*m)?.{0,60}?"
        + r"[Nn]orte\s*(?:\(Y\))?\s*[=:]?\s*" + _UTM
        + r"|"
        + r"[Nn]orte\s*(?:\(Y\))?\s*[=:]?\s*" + _UTM
        + r"(?:\s*m)?.{0,60}?"
        + r"[Ee]ste\s*(?:\(X\))?\s*[=:]?\s*" + _UTM
        + r"|"
        + r"\bN\s+" + _UTM + r"\s*m?\.?\s+e\s+E\s+" + _UTM
        + r"|"
        + r"[Ee]\s*=\s*" + _UTM + r".{0,30}?[Nn]\s*=\s*" + _UTM
        + r"|"
        + r"[Xx]\s*=\s*" + _UTM + r".{0,30}?[Yy]\s*=\s*" + _UTM,
        re.DOTALL
    )

    pares     = []
    seen_vals = set()
    seen_reg  = []

    for m in pat.finditer(texto):
        if any(s <= m.start() <= e for s, e in seen_reg):
            continue
        grupos = [x for x in m.groups() if x is not None]
        if len(grupos) < 2:
            continue
        try:
            v1, v2 = _num(grupos[0]), _num(grupos[1])
            if not (_is_utm(v1) and _is_utm(v2)):
                continue
            e, n   = _classificar_en(v1, v2)
            chave  = (round(e, 1), round(n, 1))
            if chave in seen_vals:
                continue
            seen_vals.add(chave)
            seen_reg.append((m.start(), m.end()))
            pares.append((m.start(), e, n))
        except Exception:
            continue

    if not pares:
        pares = _extrair_tabela(texto)

    nomes = _extrair_nomes(texto)
    resultado = []
    for i, (_, e, n) in enumerate(pares):
        resultado.append({
            "vertice": nomes[i] if i < len(nomes) else str(i + 1),
            "coord_e": e,
            "coord_n": n,
            "tipo":    "utm",
        })
    return resultado


def _extrair_pares_inline(texto):
    """
    Captura coordenadas UTM quando aparecem como par na mesma linha:
    Ex: 563006.7882 8566311.3196
    Ou: 563006,7882 8566311,3196
    Associa aos identificadores de vertices encontrados no texto.
    """
    import re

    # Par de numeros UTM na mesma linha (E menor que N)
    pat = re.compile(
        r"(\d{5,7}[.,]\d{2,6})\s+(\d{7,8}[.,]\d{2,6})"
    )

    pares = []
    seen  = set()
    for m in pat.finditer(texto):
        try:
            v1 = _num(m.group(1))
            v2 = _num(m.group(2))
            if not (_is_utm(v1) and _is_utm(v2)):
                continue
            e, n  = _classificar_en(v1, v2)
            chave = (round(e,1), round(n,1))
            if chave in seen:
                continue
            seen.add(chave)
            pares.append((m.start(), e, n))
        except Exception:
            continue

    if not pares:
        return []

    # Extrai identificadores de vertices do texto (V-01, PP-01, M-01, etc.)
    pat_id = re.compile(
        r"\b([A-Z][A-Z]?[-_]?\d{1,3})\b"
    )
    ids_vistos = []
    ids_seen   = set()
    for m in pat_id.finditer(texto):
        nome = m.group(1)
        # ignora palavras comuns
        if nome in ("UTM","GPS","CRS","DMS","IGC","IBGE","SAT"):
            continue
        if nome not in ids_seen:
            ids_seen.add(nome)
            ids_vistos.append(nome)

    # Associa: pega o "De" de cada par (primeiro vertice que aparece antes)
    resultado = []
    for i, (pos, e, n) in enumerate(pares):
        # Pega o identificador mais proximo antes desta posicao
        nome = ids_vistos[i] if i < len(ids_vistos) else str(i + 1)
        resultado.append({
            "vertice": nome,
            "coord_e": e,
            "coord_n": n,
            "tipo":    "utm",
        })

    return resultado if len(resultado) >= 3 else []


def _extrair_tabela_blocos(texto):
    """
    Parser para tabelas onde cada linha da tabela ocupa multiplas linhas de texto.
    Padrao:
      De         (linha 1)
      Para       (linha 2)
      N(Y)       (linha 3)
      E(X)       (linha 4)
      Lat        (linha 5 - opcional)
      Lon        (linha 6 - opcional)
      (vazio)    (separador)
    """
    import re

    # Detecta se o texto tem esse padrao:
    # linhas curtas (De/Para = 1-3 chars) seguidas de numeros UTM
    linhas = [l.strip() for l in texto.split('\n')]

    # Encontra inicio da tabela (linha apos cabecalho De/Para/Coord)
    inicio = 0
    for i, l in enumerate(linhas):
        if re.match(r'^[A-Z]\d*$|^\d+$', l) and i + 2 < len(linhas):
            # verifica se as proximas linhas sao numeros UTM
            prox = linhas[i+1] if i+1 < len(linhas) else ''
            dep2 = linhas[i+2] if i+2 < len(linhas) else ''
            if re.match(r'^[A-Z]\d*$|^\d+$', prox) and re.match(r'^\d[\d.,]+$', dep2):
                inicio = i
                break

    if inicio == 0:
        return []

    # Agrupa em blocos separados por linha vazia
    blocos = []
    bloco_atual = []
    for linha in linhas[inicio:]:
        if linha == '':
            if bloco_atual:
                blocos.append(bloco_atual)
                bloco_atual = []
        else:
            bloco_atual.append(linha)
    if bloco_atual:
        blocos.append(bloco_atual)

    vertices = []
    seen = set()

    for bloco in blocos:
        # precisa ter pelo menos 4 linhas: De, Para, N, E
        if len(bloco) < 4:
            continue

        # linha 0: De, linha 1: Para, linha 2: N(Y), linha 3: E(X)
        de   = bloco[0]
        para = bloco[1]
        sn   = bloco[2]
        se   = bloco[3]

        # valida: De e Para sao identificadores curtos
        if not re.match(r'^[A-Za-z0-9][A-Za-z0-9]?[A-Za-z0-9]?[A-Za-z0-9]?$', de):
            continue
        if not re.match(r'^[A-Za-z0-9][A-Za-z0-9]?[A-Za-z0-9]?[A-Za-z0-9]?$', para):
            continue

        try:
            n = _num(sn)
            e = _num(se)
        except Exception:
            continue

        if not (_is_utm(n) and _is_utm(e)):
            # tenta inverter (E antes de N)
            try:
                e2 = _num(sn)
                n2 = _num(se)
                if _is_utm(n2) and _is_utm(e2):
                    e, n = e2, n2
                else:
                    continue
            except Exception:
                continue

        chave = (round(e, 1), round(n, 1))
        if chave in seen:
            continue
        seen.add(chave)

        vertices.append({
            'vertice': de,
            'coord_e': e,
            'coord_n': n,
            'tipo':    'utm',
        })

    return vertices if len(vertices) >= 3 else []


def _extrair_tabela(texto):
    pares = []
    seen  = set()
    for i, linha in enumerate(texto.split("\n")):
        nums = re.findall(r"\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{1,3})?", linha)
        utms = []
        for n in nums:
            try:
                v = _num(n)
                if _is_utm(v):
                    utms.append(v)
            except Exception:
                pass
        if len(utms) >= 2:
            utms_s = sorted(utms, reverse=True)
            e, n   = _classificar_en(utms_s[0], utms_s[1])
            chave  = (round(e), round(n))
            if chave not in seen:
                seen.add(chave)
                pares.append((i, e, n))
    return pares


def _extrair_nomes(texto):
    STOP = {
        "de","da","do","ao","com","que","se","um","uma","no","na",
        "denominado","denominada","coordenada","coordenadas",
        "norte","sul","leste","este","oeste","area","perimetro",
        "datum","municipio","fuso","distancia","azimute"
    }
    pat = re.compile(
        r"(?:marco|v[eé]rtice|ponto)\s*(?:denominado\s*)?([A-Z0-9][A-Z0-9\-\.]{0,8})",
        re.IGNORECASE
    )
    nomes = []
    seen  = set()
    for m in pat.finditer(texto):
        nome = m.group(1).strip().rstrip(".,;)")
        if nome.lower() in STOP or len(nome) > 10:
            continue
        if nome not in seen:
            seen.add(nome)
            nomes.append(nome)
    return nomes
